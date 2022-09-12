from get_external_info import get_data
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT


# Настройка полей для документа, можно добавить доп настройки: размер документа, ориентация и т.д.
def set_margin(doc):
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1)
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(0.5)


# Возвращает номер этажа, на котором будет располагаться данная картинка
def get_picture_number_of_floor(picture):
    if picture[1].isdigit():
        floor_num = picture[:2]
    else:
        floor_num = picture[0]
    return floor_num


# Возвращает порядковый номер скрипта картинки для функции add_picture_caption
def get_script_number_of_picture(picture):
    script_num = picture[picture.find("fl")+3:picture.find("fl")+5]
    return script_num


# Возвращает отсортированный список этажей
def get_floor_list(pictures_list):
    floor_list = []
    for elem in pictures_list:
        floor_num = get_picture_number_of_floor(elem)
        if floor_num not in floor_list:
            floor_list.append(floor_num)
    return sorted(floor_list)


# Возвращает список картинок для конкретного этажа
def get_list_of_floor_pictures(floor, _all_pictures):
    pictures_list = []
    for elem in _all_pictures:
        pic_num = get_picture_number_of_floor(elem)
        if str(floor) == pic_num:
            pictures_list.append(elem)
    return pictures_list

# Добавляем пустую строку в файл
def add_blank(doc, number_of_blanks):
    for i in range(number_of_blanks):
        doc.add_paragraph()
# Добавляет первую страницу в отчет
def add_first_page_in_doc(doc):

    # Заголовок файла
    report_title = doc.add_paragraph()
    report_title.alignment  = WD_PARAGRAPH_ALIGNMENT.CENTER
    report_title_text = report_title.add_run("Отчет \nпо результатам проведения Indoor-измерений")
    report_title_text.font.name = 'Times new roman'
    report_title_text.font.size = Pt(20)
    report_title_text.bold = True
    add_blank(doc, 2)

    # Добавление лого
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run("")
    run.add_picture("logo.png", width=Pt(100))
    add_blank(doc, 3)

    # Данные об измерениях
    # Дата
    date = '04.05.2022'
    p = doc.add_paragraph()
    date_title = p.add_run("Дата выезда: ")
    date_title.font.name = 'Times new roman'
    date_title.font.size = Pt(14)
    date_title.bold = True
    date_text = p.add_run(f"{date}")
    date_text.font.name = 'Times new roman'
    date_text.font.size = Pt(14)

    # Имена измерителя и репортера
    measurer = 'Филимонов Н.А.'
    p = doc.add_paragraph()
    measurer_title = p.add_run(f"Измерения проводил: ")
    measurer_title.font.name = 'Times new roman'
    measurer_title.font.size = Pt(14)
    measurer_title.bold = True
    measurer_text = p.add_run(f"{measurer}")
    measurer_text.font.name = 'Times new roman'
    measurer_text.font.size = Pt(14)

    reporter = 'Филимонов Н.А.'
    p = doc.add_paragraph()
    reporter_title = p.add_run(f"Отчет подготовил: ")
    reporter_title.font.name = 'Times new roman'
    reporter_title.font.size = Pt(14)
    reporter_title.bold = True
    reporter_text = p.add_run(f"{reporter} с помощью Report Builder by Filimonov")
    reporter_text.font.name = 'Times new roman'
    reporter_text.font.size = Pt(14)
    add_blank(doc, 2)

    # Данные об объекте
    building = "БЦ 'Юникон'"
    p = doc.add_paragraph()
    building_title = p.add_run(f"Объект: ")
    building_title.font.name = 'Times new roman'
    building_title.font.size = Pt(14)
    building_title.bold = True
    building_text = p.add_run(f"{building}")
    building_text.font.name = 'Times new roman'
    building_text.font.size = Pt(14)

    full_address = 'Москва, улица Плеханова, д. 4А'
    p = doc.add_paragraph()
    full_address_title = p.add_run(f"Адрес Объекта: ")
    full_address_title.font.name = 'Times new roman'
    full_address_title.font.size = Pt(14)
    full_address_title.bold = True
    full_address_text = p.add_run(f"{full_address}")
    full_address_text.font.name = 'Times new roman'
    full_address_text.font.size = Pt(14)
    add_blank(doc, 1)

    # Цель работы
    doc.add_paragraph("_" * 132)
    add_blank(doc, 1)
    purpose_of_work = 'Выполнение объективной оценки технических возможностей сети ПАО «Мегафон» на территории ' \
                      'объекта, после установки на нем indoor базовой станции ПАО «МегаФон» БС'
    site_id = 25456
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(21)
    purpose_of_work_title = p.add_run("        Цель работы: ")
    purpose_of_work_title.font.name = 'Times new roman'
    purpose_of_work_title.font.size = Pt(14)
    purpose_of_work_title.bold = True
    purpose_of_work_text = p.add_run(f"{purpose_of_work} {site_id}.")
    purpose_of_work_text.font.name = 'Times new roman'
    purpose_of_work_text.font.size = Pt(14)

    doc.add_page_break()


# Добавление нижнего колонтитула в файл, на вход принимает документ
def add_footer_in_doc(doc):
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]

    # Должен вызывать функцию для ввода информации пользователем {site_ID} {location} {full_address}
    site_id = '25654'
    location = 'БЦ Юникон'
    full_address = 'Москва, Плеханова, д.4А'
    footer_text = footer_para.add_run(f'Indoor, SiteID - {site_id}\n{location}\n{full_address}')
    footer_text.font.name = 'Times new roman'
    footer_text.font.size = Pt(8)


# Добавляет верхний колонтитул
def add_header_in_doc(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    header_text = header_para.add_run(f'Отчет по результатам проведения\n'
                                      f'Indoor-измерений в сети ПАО «МегаФон»')

    header_text.font.name = 'Times new roman'
    header_text.font.size = Pt(11)
    header_text.italic = True


# Подготавливает дату для таблицы этажа
def get_data_for_one_floor(_path, floor):
    data = get_data(_path)
    i = 0
    data_for_one_floor = []
    for data[i] in data:

        # Костыль т.к. если этаж будет написан fl10 - Не сработает
        if data[i].find(str(floor) + 'fl') != -1:
            for j in range(0,4):
                data_for_one_floor.append(data[i+j])
        i += 1
    return data_for_one_floor

# Добавляет заголовок этажа и форматирует его
def add_floor_title_in_file(floor, doc):
    floor_title = doc.add_heading().add_run(f'{floor} этаж')
    floor_title.font.name = 'Times new roman'
    floor_title.font.size = Pt(16)
    floor_title.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()


# Добавляет таблицу для заполнения значениями
def add_table_with_values(doc, data):
    rows = 4
    cols = 4
    table = doc.add_table(rows=rows, cols=cols)
    table.rows[0].cells[1].text = '2G'
    table.rows[0].cells[2].text = '3G'
    table.rows[0].cells[3].text = '4G'
    table.rows[1].cells[0].text = 'Средний уровень сигнала, дБм'
    table.rows[2].cells[0].text = 'Cредняя скорость DL, Мб/c'
    table.rows[3].cells[0].text = 'Cредняя скорость UL, Мб/c'
    table.rows[1].cells[1].text = '-'  # Rx level
    table.rows[2].cells[1].text = '-'  # DL 2G
    table.rows[3].cells[1].text = '-'  # UL 2G
    table.rows[1].cells[2].text = 'RSCP'  # RSCP
    table.rows[2].cells[2].text = 'DL 3G'  # DL 3G
    table.rows[3].cells[2].text = 'UL 3G'  # UL 3G
    table.rows[1].cells[3].text = 'RSRP'  # RSRP
    table.rows[2].cells[3].text = 'DL 4G'  # DL 2G
    table.rows[3].cells[3].text = 'UL 4G'  # UL 2G

    doc.add_paragraph()
    entering_data_into_a_table(table, data)
    adjust_table_with_values(table, rows, cols)


# Наполнение таблицы данными из файла
def entering_data_into_a_table(table, data):

    for i in range(len(data)-3):

        if data[i][-2:] == ':7':
            table.rows[1].cells[1].text = data[i+1]             # Rx level
            if data[i+2] != '-1.0':
                table.rows[2].cells[1].text = data[i+2]         # DL 2G
            if data[i+3] != '-1.0':
                table.rows[3].cells[1].text = data[i+3]         # UL 2G

        if  data[i][-2:] == ':1':
            table.rows[1].cells[2].text = data[i+1]             # RSCP
            if data[i + 2] != '-1.0':
                table.rows[2].cells[2].text = data[i + 2]       # DL 3G
            if data[i + 3] != '-1.0':
                table.rows[3].cells[2].text = data[i + 3]       # UL 3G

        if data[i][-2:] == ':3':
            table.rows[1].cells[3].text = data[i + 1]           # RSRP
            if data[i + 2] != '-1.0':
                table.rows[2].cells[3].text = data[i + 2]       # DL 4G
            if data[i + 3] != '-1.0':
                table.rows[3].cells[3].text = data[i + 3]       # UL 4G


# Добавляет 1 картинку файл
def add_picture_in_file(_path, doc, picture):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run("")
    run.add_picture(_path + '/' + picture, width=Pt(460))


# Форматирование таблицы со значениями
def adjust_table_with_values(tab, rows, cols):
    tab.alignment = WD_TAB_ALIGNMENT.CENTER
    for cell in tab.columns[0].cells:
        cell.width = Cm(7.5)
    for i in range(rows):
        for j in range(1, cols):
            tab.cell(i, j).paragraphs[0].paragraph_format.alignment = WD_TAB_ALIGNMENT.CENTER


# Добавление названия к картинке
def add_picture_caption(picture, doc, serial_number):
    picture_captions = {'01': 'Качественные показатели покрытия 2G', '02': 'Качественные показатели покрытия 3G',
                        '03': 'Качественные показатели покрытия 4G', '04': 'Скорость ППД DL/UL 3G',
                        '05': 'Скорость ППД DL/UL 4G', '06': 'Функциональные показатели CSFB',
                        '07': 'Функциональные показатели LTE Carrier Aggregation',
                        '08': 'Функциональные показатели LTE MIMO', '09': 'Функциональные показатели 2G indoor scanner',
                        '10': 'Функциональные показатели 3G indoor scanner',
                        '11': 'Функциональные показатели LTE indoor scanner'
                        }
    script_number_of_picture = get_script_number_of_picture(picture)
    picture_caption = picture_captions.get(script_number_of_picture)
    paragraph = doc.add_paragraph('')
    text = paragraph.add_run(f'Рисунок {serial_number} - {picture_caption}')

    # Форматирование подписи
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    text.font.name = 'Times new roman'
    text.font.size = Pt(14)


# Добавление подписи для раздела измерений со сканера
def add_scanner_title(doc):
    doc.add_page_break()
    floor_title = doc.add_heading().add_run('Функциональные показатели со сканера')
    floor_title.font.name = 'Times new roman'
    floor_title.font.size = Pt(14)
    floor_title.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()


# Полная "Сборка" 1 этажа для отчета
def add_floor_in_report(_path, doc, floor, floor_pictures_list, picture_number_in_file):
    add_floor_title_in_file(floor, doc)
    data = get_data_for_one_floor(_path, floor)
    add_table_with_values(doc, data)
    is_scanner_title_was_add = 0
    for picture in floor_pictures_list:
        script_number = get_script_number_of_picture(picture)
        if script_number in ['09', '10', '11'] and is_scanner_title_was_add == 0:
            add_scanner_title(doc)
            is_scanner_title_was_add = 1
        add_picture_in_file(_path, doc, picture)
        add_picture_caption(picture, doc, picture_number_in_file)
        picture_number_in_file += 1
    doc.add_page_break()
    return picture_number_in_file
