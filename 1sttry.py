import os
from tkinter import filedialog
from docx import Document
from docx.shared import Pt, Cm



def select_path():
    path = ''
    while path == '':
        path = filedialog.askdirectory(title="Select a File")
    return path


def get_pictures(path):
    all_files = os.listdir(path)
    list_of_pictures = []
    for file in all_files:
        if file[-4:] == '.jpg' or file[-4:] == '.png':
            list_of_pictures.append(file)
    return list_of_pictures


def add_pictures_in_file(list_of_pictures, doc):
    for picture in list_of_pictures:
        doc.add_picture(path_with_pictures + '/' + picture, width=Pt(500))


def margin_set(doc):
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1)


path = 'C:/Users/PC/Desktop/Work/'  #legacy for save time
report_doc = Document()
margin_set(report_doc)
path_with_pictures = path + '/Pictures'
#path_with_pictures = select_path()
all_pictures = get_pictures(path_with_pictures)
add_pictures_in_file(all_pictures, report_doc)
#report_path = select_path()
report_path = path
try:
    report_doc.save(report_path + '/test.docx')
    print('SUCCESS\n')
except PermissionError:
    print('Close the file pls\n')
