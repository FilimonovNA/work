import os
from tkinter import filedialog
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT

'''Добавить функции вычисления номера этажа и номера скрипта картинки'''


def select_path():
    user_path = ''
    while user_path == '':
        user_path = filedialog.askdirectory(title="Select a File")
    return user_path


def get_pictures_list(user_path):
    all_files = os.listdir(user_path)
    list_of_pictures = []
    for file in all_files:
        if file[-4:] == '.jpg' or file[-4:] == '.png':
            list_of_pictures.append(file)
    return list_of_pictures


def add_picture_in_file(picture, doc):
    doc.add_picture(path_with_pictures + '/' + picture, width=Pt(480))


def set_margin(doc):
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1)


def get_picture_number_of_floor(picture):
    if picture[1].isdigit():
        floor_num = picture[:2]
    else:
        floor_num = picture[0]
    return floor_num


def get_floor_list(pictures_list):
    elem = ''
    floor_list = []
    for elem in pictures_list:
        floor_num = get_picture_number_of_floor(elem)
        if floor_num not in floor_list:
            floor_list.append(floor_num)
    return sorted(floor_list)


def add_floor_title_in_file(floor, doc):
    floor_title = doc.add_heading().add_run(f'{floor} этаж')
    floor_title.font.name = 'Times new roman'
    floor_title.font.size = Pt(16)
    floor_title.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()


def floor_pictures(current_floor, _all_pictures):
    pictures_list = []
    for elem in _all_pictures:
        pic_num = get_picture_number_of_floor(elem)
        if str(current_floor) == pic_num:
            pictures_list.append(elem)
    return pictures_list


def table_adjusting(tab):
    for cell in tab.columns[0].cells:
        cell.width = Cm(7.5)
    for col in tab.columns:
        for cell in col.cells:
            cell.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_table_with_values(doc):
    rows = 4
    cols = 4
    table = doc.add_table(rows=rows, cols=cols)
    table.rows[0].cells[1].text = '2G'
    table.rows[0].cells[2].text = '3G'
    table.rows[0].cells[3].text = '4G'

    table.rows[1].cells[0].text = 'Средний уровень сигнала, дБм'
    table.rows[2].cells[0].text = 'Cредняя скорость DL, Мб/c'
    table.rows[3].cells[0].text = 'Cредняя скорость UL, Мб/c'
    table.rows[1].cells[1].text = 'Meas Rx Level'           # Rx level
    table.rows[2].cells[1].text = '-'                       # DL 2G
    table.rows[3].cells[1].text = '-'                       # UL 2G
    table.rows[1].cells[2].text = 'RSCP'                    # RSCP
    table.rows[2].cells[2].text = 'DL3G'                    # DL 3G
    table.rows[3].cells[2].text = 'UL3G'                    # UL 3G
    table.rows[1].cells[3].text = 'RSRP'                    # RSRP
    table.rows[2].cells[3].text = 'DL4G'                    # DL 2G
    table.rows[3].cells[3].text = 'UL4G'                    # UL 2G
    table_adjusting(table)


def add_picture_caption(picture, doc, serial_number):
    picture_captions = {
    '01': 'Качественные показатели покрытия 2G', '02': 'Качественные показатели покрытия 3G',
    '03': 'Качественные показатели покрытия 4G', '04': 'Скорость ППД DL/UL 3G',
    '05': 'Скорость ППД DL/UL 4G', '06': 'Функциональные показатели CSFB',
    '07': 'Функциональные показатели LTE Carrier Aggregation',
    '08': 'Функциональные показатели LTE MIMO', '09': 'Функциональные показатели 2G indoor',
    '10': 'Функциональные показатели 3G indoor', '11': 'Функциональные показатели LTE indoor'
    }
    picture_number_on_floor = picture[picture.find("fl")+3:picture.find("fl")+5]
    picture_caption = picture_captions.get(picture_number_on_floor)
    paragraph = doc.add_paragraph('')
    text = paragraph.add_run(f'Рисунок {serial_number} - {picture_caption}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    text.font.name = 'Times new roman'
    text.font.size = Pt(14)


'''Если номер картинки == 9, 10, 11 - выводим add scanner title'''


def add_scanner_title(doc):
    floor_title = doc.add_heading().add_run('Функциональные показатели со сканера')
    floor_title.font.name = 'Times new roman'
    floor_title.font.size = Pt(14)
    floor_title.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()


def construct_floor(report_doc, floor, floor_pictures_list, absolut_picture_number_in_file):
    add_floor_title_in_file(floor, report_doc)
    add_table_with_values(report_doc)
    for picture in floor_pictures_list:
        add_picture_in_file(picture, report_doc)
        add_picture_caption(picture, report_doc, absolut_picture_number_in_file)
        absolut_picture_number_in_file += 1
    report_doc.add_page_break()
    return absolut_picture_number_in_file


path = 'C:/Users/PC/Desktop/Work/'  # legacy for save time
report_doc = Document()
set_margin(report_doc)
path_with_pictures = path + '/Pictures'
# path_with_pictures = select_path()
all_pictures = get_pictures_list(path_with_pictures)
all_floors = get_floor_list(all_pictures)


absolut_picture_number_in_file = 1
for floor in all_floors:
    floor_pictures_list = floor_pictures(floor, all_pictures)  # получаем список картинок для конкретного этажа
    absolut_picture_number_in_file = construct_floor(report_doc, floor, floor_pictures_list, absolut_picture_number_in_file)  # создаем конструкцию первого этажа
# report_path = select_path()
report_path = path

try:
    report_doc.save(report_path + '/test.docx')
    print('SUCCESS')
except PermissionError:
    print('Close the file pls')
