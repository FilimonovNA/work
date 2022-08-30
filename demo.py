'''
        Просматриваем папку с файлами,
        ставим в соответствие номер этажа и названия картинок/скриптов для этажа
'''
from tkinter import filedialog
import os
import docx
from docx import Document

def select_file():
   path= filedialog.askdirectory(title="Select a File")
   return path

path = select_file()
print(path)

pictures_directory = path
all_pictures = os.listdir(pictures_directory)
floors = []
full_floors_pictures = {}
full_floors_scripts = {}
#создаем список этажей
for picture in all_pictures:
    floor_name = picture[:2]
    if floor_name in floors:
        pass
    else:
        floors.append(picture[:2])
#словарь с этажами и скриптами
for floor_number in floors:
    floor_scripts = []
    floor_pictures = []
    for picture in all_pictures:
        picture_floor = picture[:2]
        if picture_floor == floor_number:
            floor_scripts.append(picture[5:-3])
            floor_pictures.append(picture)
    full_floors_pictures[floor_number] = floor_pictures
    full_floors_scripts[floor_number] = floor_scripts
#print(full_floors_pictures, '\n', full_floors_scripts)
'''
        Далее создаем документ и редактируем его
'''
#Добавляем картинки
doc = Document() #создаем пустой файл
file_path = f'c:\\result\\demo.docx'



for floor in floors:
    floor_title = doc.add_paragraph().add_run(f'{floor} этаж')
    # размер шрифта
    floor_title.font.name = 'Times new roman'
    floor_title.font.size = Pt(16)
    floor_title.bold = True
    floor_title.color = 'Red'
    doc.add_paragraph()
    for picture in full_floors_pictures[floor]:
        doc.add_picture(f'{pictures_directory}\\{picture}', width=Pt(600))
    doc.add_page_break()

''' Настройка полей'''
sections = doc.sections
for section in sections:
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

'''
Если файл открыт нужно его закрыть, иначе ничего не запишется, кек,
- Зачем это нужно?
- Чтобы пока я пишу код не видеть кучу ошибок'''
try:
    doc.save(file_path)
    print('С кайфом все сделал\n'*10)
except PermissionError:
    print('Close the file pls\n'*10)